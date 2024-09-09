from docx import Document
import pandas as pd
from tkinter import Tk, filedialog, messagebox, Button, Label, Radiobutton, StringVar, OptionMenu

def write_multiple_docx(data, index, save_path):
    doc = Document()
    for key, value in data.items():
        doc.add_paragraph(f"{key}:{value}")
    name = f"Document{index + 1}.docx"
    doc.save(f"{save_path}/{name}")
    print(f"Document {name} created successfully")

def write_single_docx(data, doc):
    for key, value in data.items():
        doc.add_paragraph(f"{key}:{value}")
    doc.add_paragraph()
    return doc

def process_files(path, save_path, option):
    if option == 'Multiple Documents':
        d = pd.read_excel(path)
        for index, row in d.iterrows():
            data = row.to_dict()
            write_multiple_docx(data, index, save_path)
    elif option == 'Single Document':
        single_doc = Document()
        d = pd.read_excel(path)
        for index, row in d.iterrows():
            data = row.to_dict()
            single_doc = write_single_docx(data, single_doc)
        single_doc.save(f"{save_path}/SingleDocument.docx")
        print("Single Document created successfully")

def select_file():
    file_path = filedialog.askopenfilename(
        title="Select an Excel file",
        filetypes=[("Excel files", "*.xlsx")]
    )
    if file_path:
        file_var.set(file_path)

def select_folder():
    folder_path = filedialog.askdirectory(title="Select the folder to save the documents")
    if folder_path:
        folder_var.set(folder_path)

def process():
    file_path = file_var.get()
    print(file_path)
    save_path = folder_var.get()
    print(save_path)
    option = option_var.get()
    if file_path and save_path and option:
        process_files(file_path, save_path, option)
    else:
        messagebox.showwarning("Warning", "Please complete all selections.")

def main():
    global file_var, folder_var, option_var

    root = Tk()
    root.title("Document Generator")

    file_var = StringVar()
    folder_var = StringVar()
    option_var = StringVar(value='Multiple Documents')

    Label(root, text="Select Excel File:").pack()
    Button(root, text="Browse", command=select_file).pack()
    Label(root, textvariable=file_var).pack()

    Label(root, text="Select Save Folder:").pack()
    Button(root, text="Browse", command=select_folder).pack()
    Label(root, textvariable=folder_var).pack()

    Label(root, text="Select Document Type:").pack()
    Radiobutton(root, text="Multiple Documents", variable=option_var, value='Multiple Documents').pack()
    Radiobutton(root, text="Single Document", variable=option_var, value='Single Document').pack()

    Button(root, text="Generate Documents", command=process).pack()

    root.mainloop()

if __name__ == "__main__":
    main()